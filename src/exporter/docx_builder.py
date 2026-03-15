import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class DocxBuilder:
    def __init__(self, output_path: str = "../output/Summary_Report.docx"):
        self.output_path = output_path
        self.doc = Document()
        # 设置默认字体
        style = self.doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(12)

    def _add_markdown_content(self, md_text: str):
        """简单的 Markdown 解析并渲染到 docx"""
        lines = md_text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue

            # 处理标题
            if line.startswith('### '):
                self.doc.add_heading(line.replace('### ', ''), level=3)
            elif line.startswith('## '):
                self.doc.add_heading(line.replace('## ', ''), level=2)
            elif line.startswith('# '):
                self.doc.add_heading(line.replace('# ', ''), level=1)
            # 处理列表
            elif line.startswith('- ') or line.startswith('* '):
                self.doc.add_paragraph(line[2:], style='List Bullet')
            # 处理普通文本
            else:
                self.doc.add_paragraph(line)

    def save_report(self, md_content: str):
        # 移除可能存在的 Markdown 代码块标记
        clean_content = re.sub(r'```markdown|```', '', md_content).strip()
        
        self._add_markdown_content(clean_content)
        
        # 确保输出目录存在
        os.makedirs(os.path.dirname(self.output_path), exist_ok=True)
        self.doc.save(self.output_path)
        print(f"✅ 文档已生成: {self.output_path}")
        return self.output_path

if __name__ == "__main__":
    builder = DocxBuilder("../output/test_report.docx")
    sample_md = "# 一、任务概况\n## (一) 核心目标\n- 目标1\n- 目标2\n### 1. 详细细节\n普通文本内容。"
    builder.save_report(sample_md)
